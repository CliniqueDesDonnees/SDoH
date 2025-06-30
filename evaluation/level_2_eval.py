import re
import json
import unidecode
from pathlib import Path
from collections import Counter
from typing import Dict, List, Tuple, Optional
import pandas as pd
from docx import Document

### ----------------------------------------------------------
### Utility Functions
### ----------------------------------------------------------

def remove_prefix(text: str, prefix: str) -> str:
    return text[len(prefix):] if text.startswith(prefix) else text

def has_overlap(start1: int, end1: int, start2: int, end2: int) -> bool:
    return max(start1, start2) < min(end1, end2)

def get_overlap(start1: int, end1: int, start2: int, end2: int) -> List[int]:
    return sorted(set(range(start1, end1)) & set(range(start2, end2)))

def find_closest_span(entity_offset: Tuple[int, int], relation_offsets: List[Tuple[int, int]]) -> Optional[Tuple[int, int]]:
    entity_mid = (entity_offset[0] + entity_offset[1]) / 2
    return min(relation_offsets, key=lambda span: abs((span[0] + span[1]) / 2 - entity_mid), default=None)

def calc_precision(tp: int, np: int) -> float:
    return tp / np if np>0 else 0.0

def calc_recall(tp: int, nt: int) -> float:
    return tp / nt if nt>0 else 0.0

def calc_f1(p: float, r: float) -> float:
    return 2 * p * r / (p + r) if p*r!=0 else 0.0


### ----------------------------------------------------------
### SDOH Parsing and Structuring
### ----------------------------------------------------------

def parse_brat_file(txt_file: Path, annotation_file_suffixes: List[str] = None, parse_notes: bool = False) -> Dict:

    example = {}
    example["document_id"] = txt_file.with_suffix("").name
    with txt_file.open() as f:
        example["text"] = f.read()

    # If no specific suffixes of the to-be-read annotation files are given - take standard suffixes
    # for event extraction
    if annotation_file_suffixes is None:
        annotation_file_suffixes = [".a1", ".a2", ".ann"]

    if len(annotation_file_suffixes) == 0:
        raise AssertionError(
            "At least one suffix for the to-be-read annotation files should be given!"
        )

    ann_lines = []
    for suffix in annotation_file_suffixes:
        annotation_file = txt_file.with_suffix(suffix)
        if annotation_file.exists():
            with annotation_file.open() as f:
                ann_lines.extend(f.readlines())

    example["text_bound_annotations"] = []
    example["events"] = []
    example["relations"] = []
    example["equivalences"] = []
    example["attributes"] = []
    example["normalizations"] = []

    if parse_notes:
        example["notes"] = []

    for line in ann_lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("T"):  # Text bound
            ann = {}
            fields = line.split("\t")

            ann["id"] = fields[0]
            ann["type"] = fields[1].split()[0]
            ann["offsets"] = []
            span_str = remove_prefix(fields[1], (ann["type"] + " "))
            text = fields[2]
            for span in span_str.split(";"):
                start, end = span.split()
                ann["offsets"].append([int(start), int(end)])

            # Heuristically split text of discontiguous entities into chunks
            ann["text"] = []
            if len(ann["offsets"]) > 1:
                i = 0
                for start, end in ann["offsets"]:
                    chunk_len = end - start
                    ann["text"].append(text[i : chunk_len + i])
                    i += chunk_len
                    while i < len(text) and text[i] == " ":
                        i += 1
            else:
                ann["text"] = [text]

            example["text_bound_annotations"].append(ann)

        elif line.startswith("E"):
            ann = {}
            fields = line.split("\t")

            ann["id"] = fields[0]

            ann["type"], ann["trigger"] = fields[1].split()[0].split(":")

            ann["arguments"] = []
            for role_ref_id in fields[1].split()[1:]:
                argument = {
                    "role": (role_ref_id.split(":"))[0],
                    "ref_id": (role_ref_id.split(":"))[1],
                }
                ann["arguments"].append(argument)

            example["events"].append(ann)

        elif line.startswith("R"):
            ann = {}
            fields = line.split("\t")

            ann["id"] = fields[0]
            ann["type"] = fields[1].split()[0]

            ann["head"] = {
                "role": fields[1].split()[1].split(":")[0],
                "ref_id": fields[1].split()[1].split(":")[1],
            }
            ann["tail"] = {
                "role": fields[1].split()[2].split(":")[0],
                "ref_id": fields[1].split()[2].split(":")[1],
            }

            example["relations"].append(ann)

        # '*' seems to be the legacy way to mark equivalences,
        # but I couldn't find any info on the current way
        # this might have to be adapted dependent on the brat version
        # of the annotation
        elif line.startswith("*"):
            ann = {}
            fields = line.split("\t")

            ann["id"] = fields[0]
            ann["ref_ids"] = fields[1].split()[1:]

            example["equivalences"].append(ann)

        elif line.startswith("A") or line.startswith("M"):
            ann = {}
            fields = line.split("\t")

            ann["id"] = fields[0]

            info = fields[1].split()
            ann["type"] = info[0]
            ann["ref_id"] = info[1]

            if len(info) > 2:
                ann["value"] = info[2]
            else:
                ann["value"] = ""

            example["attributes"].append(ann)

        elif line.startswith("N"):
            ann = {}
            fields = line.split("\t")

            ann["id"] = fields[0]
            ann["text"] = fields[2]

            info = fields[1].split()

            ann["type"] = info[0]
            ann["ref_id"] = info[1]
            ann["resource_name"] = info[2].split(":")[0]
            ann["cuid"] = info[2].split(":")[1]
            example["normalizations"].append(ann)

        elif parse_notes and line.startswith("#"):
            ann = {}
            fields = line.split("\t")

            ann["id"] = fields[0]
            ann["text"] = fields[2] if len(fields) == 3 else "<BB_NULL_STR>"

            info = fields[1].split()

            ann["type"] = info[0]
            ann["ref_id"] = info[1]
            example["notes"].append(ann)
    return example

def convert_text_to_struct(example: Dict) -> Dict:

    main_entities = ["Living_Alone", "Living_WithOthers", "MaritalStatus_Single", "MaritalStatus_InRelationship", "MaritalStatus_Divorced", "MaritalStatus_Widowed", "Descendants_Yes", "Descendants_No", "Job", "Last_job", "Employment_Working", "Employment_Unemployed", "Employment_Student", "Employment_Pensioner", "Employment_Other", "Housing_Yes", "Housing_No", "PhysicalActivity_Yes", "PhysicalActivity_No", "Income", "Education", "Ethnicity"]
    relation_entities = ["StatusTime", "History", "Duration", "Amount", "Frequency", "Type"]
    event_entities = ["Tobacco", "Alcohol", "Drug"]

    output_sdoh = []

    list_entities = {m['id']: m for m in example["text_bound_annotations"]}
    list_attributes = {m['ref_id']: m for m in example["attributes"]}

    for i in example["text_bound_annotations"]:

        if i['type'] in main_entities:

            res_entity = (i['type'], i['text'][0], i['offsets'][0])
            res_relations = []

            id_entity = i['id']
            for r in example["relations"]:
                if id_entity == r['head']['ref_id']:
                    entity_rel = r['tail']['ref_id']
                    res_relations.append({
                        "relation": list_entities[entity_rel]['type'],
                        "text": list_entities[entity_rel]['text'][0],
                        "offsets": list_entities[entity_rel]['offsets'][0]
                        })
            output_sdoh.append({
                "entity": res_entity,
                "relations": res_relations
                })
        elif i['type'] in event_entities:
            res_entity = (i['type'], i['text'][0], i['offsets'][0])
            res_relations = []

            id_entity = i['id']
            for r in example["relations"]:
                if id_entity == r['head']['ref_id']:
                    entity_rel = r['tail']['ref_id']
                    if list_entities[entity_rel]['type'] == "StatusTime":
                        res_relations.append({
                            "relation": list_entities[entity_rel]['type'],
                            "attribute": list_attributes[entity_rel]['value'],
                            "text": list_entities[entity_rel]['text'][0],
                            "offsets": list_entities[entity_rel]['offsets'][0]
                            })
                    else:
                        res_relations.append({
                            "relation": list_entities[entity_rel]['type'],
                            "text": list_entities[entity_rel]['text'][0],
                            "offsets": list_entities[entity_rel]['offsets'][0]
                            })
            output_sdoh.append({
                "entity": res_entity,
                "relations": res_relations
                })
    return output_sdoh

def parse_structured_sdoh(text: str) -> List[Dict]:
    if not text:
        return [{"entity": None, "relations": []}]

    entities = []
    # Extract entities (between « and »)
    pattern_entity = r"«(.*?)» (.*?)(?= \[|$)"
    # Extract relations (between [] and additional attributes if applicable)
    pattern_relation = r"\[(.*?)\](?:\[(.*?)\])? (.*?)(?=\[|$)"

    entity_match = re.search(pattern_entity, text)
    if entity_match:
        entities.append((entity_match.group(1), entity_match.group(2).strip()))

    relations = []
    for match in re.finditer(pattern_relation, text):
        relation = match.group(1)
        attribute = match.group(2)
        rel_text = match.group(3).strip()
        rel_entry = {"relation": relation, "text": rel_text}
        if attribute:
            rel_entry["attribute"] = attribute
        relations.append(rel_entry)

    return [{"entity": entities[0] if entities else None, "relations": relations}]


def format_struct_output(text: str) -> List[Dict]:
    return [parse_structured_sdoh(entry)[0] for entry in text.split(" €AND ")]


def update_entity_offsets(entry: Dict, full_text: str, previous_entities: List[Dict]) -> Optional[Dict]:
    text = unidecode.unidecode(full_text.lower())
    if not entry.get("entity"):
        return None

    label, value = entry["entity"][:2]
    pattern = re.compile(re.escape(unidecode.unidecode(value.lower())))
    offsets = [(m.start(), m.end()) for m in pattern.finditer(text)]

    prior = [e for e in previous_entities if e['entity'][0] == label]

    if not offsets:
        return None
    if len(offsets) == 1:
        entry["entity"] += (list(offsets[0]),)
    else:
        non_overlapping = [o for o in offsets if not any(has_overlap(o[0], o[1], e["entity"][2][0], e["entity"][2][1]) for e in prior)]
        if non_overlapping:
            entry["entity"] += (list(non_overlapping[0]),)
        else:
            entry["entity"] += (list(offsets[0]),)
    return entry


def update_relation_offsets(entry: Dict, full_text: str) -> Optional[Dict]:
    if not entry.get("entity"):
        return None
    text = unidecode.unidecode(full_text.lower())
    updated = []

    for rel in entry.get("relations", []):
        pattern = re.compile(re.escape(unidecode.unidecode(rel["text"].lower())))
        matches = [(m.start(), m.end()) for m in pattern.finditer(text)]

        if not matches:
            continue

        if len(matches) == 1:
            rel["offsets"] = list(matches[0])
        else:
            closest = find_closest_span(entry["entity"][2], matches)
            if closest:
                rel["offsets"] = list(closest)
        updated.append(rel)

    entry["relations"] = updated
    return entry


def update_status_relations(entry: Dict) -> Dict:
    for rel in entry.get("relations", []):
        if ":" in rel["relation"]:
            rel_type, attr = rel["relation"].split(":")
            rel["relation"] = rel_type
            rel["attribute"] = attr
    return entry


### ----------------------------------------------------------
### Evaluation Functions
### ----------------------------------------------------------

def count_events(events: List[Dict], key: Tuple[str, str, str]) -> int:
    count = 0
    for ev in events:
        entity, rels = ev.get("entity"), ev.get("relations", [])
        if not entity:
            continue
        ent_type = entity[0]
        if key == (ent_type, "N/A", "N/A"):
            count += 1
        for r in rels:
            r_type = r.get("relation")
            attr = r.get("attribute", "N/A")
            if key == (ent_type, r_type, attr):
                count += 1
    return count


def filter_events(events: List[Dict], key: Tuple[str, str, str]) -> List[Dict]:
    filtered = []
    for ev in events:
        entity, rels = ev.get("entity"), ev.get("relations", [])
        if not entity:
            continue
        ent_type = entity[0]
        if key == (ent_type, "N/A", "N/A"):
            filtered.append(ev)
        for r in rels:
            r_type = r.get("relation")
            attr = r.get("attribute", "N/A")
            if key == (ent_type, r_type, attr):
                filtered.append({"entity": entity, "relations": r})
    return filtered


def match_entities(gold: List[Dict], pred: List[Dict], key: Tuple[str, str, str], match_type: str) -> int:
    matched_gold, matched_pred = set(), set()
    matches = 0

    for i, g in enumerate(gold):
        for j, p in enumerate(pred):
            if i in matched_gold or j in matched_pred:
                continue

            g_ent = g['entity'][2]
            p_ent = p['entity'][2]

            g_span_match = g_ent == p_ent if match_type == "exact" else has_overlap(*g_ent, *p_ent)

            if key[1:] == ("N/A", "N/A"):
                if g_span_match:
                    matches += 1
                    matched_gold.add(i)
                    matched_pred.add(j)
            else:
                g_rel = g['relations']['offsets']
                p_rel = p['relations']['offsets']
                rel_match = g_rel == p_rel if match_type == "exact" else has_overlap(*g_rel, *p_rel)
                if g_span_match and rel_match:
                    matches += 1
                    matched_gold.add(i)
                    matched_pred.add(j)

    return matches


def get_event_matches(gold: List[Dict], pred: List[Dict], key: Tuple[str, str, str], match_type: str) -> int:
    g_filtered = filter_events(gold, key)
    p_filtered = filter_events(pred, key)
    if not p_filtered:
        return 0
    return match_entities(g_filtered, p_filtered, key, match_type)


### ----------------------------------------------------------
### Output and Export
### ----------------------------------------------------------

def build_eval_dataframe(results: List[Dict], combinations: List[Tuple[str, str, str]], eval_type: str = "exact") -> pd.DataFrame:
    records = []

    for entity, relation, argument in combinations:
        NT = NP = TP = 0

        for entry in results:
            gold = entry.get('gold_format_updated', [])
            pred = entry.get('pred_format_updated', [])

            NT += count_events(gold, (entity, relation, argument))
            NP += count_events(pred, (entity, relation, argument))
            TP += get_event_matches(gold, pred, (entity, relation, argument), eval_type)

        P = calc_precision(TP, NP)
        R = calc_recall(TP, NT)
        F1 = calc_f1(P, R)

        records.append({
            "entity": entity,
            "relation": relation,
            "argument": argument,
            "NT": NT,
            "NP": NP,
            "TP": TP,
            "P": round(P, 4),
            "R": round(R, 4),
            "F1": round(F1, 4)
        })

    return pd.DataFrame.from_records(records)


def save_df_to_word(df: pd.DataFrame, output_path: str) -> None:
    doc = Document()
    doc.add_heading('Evaluation Report', level=1)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    for i, column_name in enumerate(df.columns):
        table.cell(0, i).text = column_name

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    doc.save(output_path)
    print(f"Saved evaluation results to {output_path}")


### ----------------------------------------------------------
### Main Execution Pipeline
### ----------------------------------------------------------

def main():
    import argparse

    parser = argparse.ArgumentParser(description="Evaluate SDOH predictions against gold annotations.")
    parser.add_argument('--input', type=str, required=True, help='Path to input JSONL file with predictions')
    parser.add_argument('--brat-dir', type=str, required=True, help='Directory containing gold BRAT annotations')
    parser.add_argument('--output', type=str, default='sdoh_eval_results.docx', help='Output Word file path')
    parser.add_argument('--eval-type', type=str, default='exact', choices=['exact', 'overlap'], help='Evaluation span match type')
    args = parser.parse_args()

    # Load input predictions
    with open(args.input, 'r', encoding='utf-8') as f:
        json_entries = [json.loads(line) for line in f]

    # Define all possible combinations of entities, relations, attributes
    possible_combinations = [
        ('Job', 'N/A', 'N/A'),
        ('Last_job', 'N/A', 'N/A'),
        ('Descendants_Yes', 'N/A', 'N/A'),
        ('Descendants_Yes', 'Type', 'N/A'),
        ('Descendants_Yes', 'Amount', 'N/A'),
        ('Descendants_No', 'N/A', 'N/A'),
        ('Housing_Yes', 'N/A', 'N/A'),
        ('Housing_Yes', 'History', 'N/A'),
        ('Housing_No', 'N/A', 'N/A'),
        ('MaritalStatus_InRelationship', 'N/A', 'N/A'),
        ('MaritalStatus_InRelationship', 'History', 'N/A'),
        ('MaritalStatus_Divorced', 'N/A', 'N/A'),
        ('MaritalStatus_Divorced', 'History', 'N/A'),
        ('MaritalStatus_Single', 'N/A', 'N/A'),
        ('MaritalStatus_Widowed', 'N/A', 'N/A'),
        ('MaritalStatus_Widowed', 'History', 'N/A'),
        ('Living_Alone', 'N/A', 'N/A'),
        ('Living_WithOthers', 'N/A', 'N/A'),
        ('Living_WithOthers', 'History', 'N/A'),
        ('Employment_Working', 'N/A', 'N/A'),
        ('Employment_Working', 'History', 'N/A'),
        ('Employment_Other', 'N/A', 'N/A'),
        ('Employment_Other', 'History', 'N/A'),
        ('Employment_Pensioner', 'N/A', 'N/A'),
        ('Employment_Pensioner', 'History', 'N/A'),
        ('Employment_Unemployed', 'N/A', 'N/A'),
        ('Employment_Unemployed', 'History', 'N/A'),
        ('Employment_Student', 'N/A', 'N/A'),
        ('Alcohol', 'StatusTime', 'current'),
        ('Alcohol', 'StatusTime', 'none'),
        ('Alcohol', 'StatusTime', 'past'),
        ('Alcohol', 'Type', 'N/A'),
        ('Alcohol', 'Frequency', 'N/A'),
        ('Alcohol', 'Amount', 'N/A'),
        ('Alcohol', 'History', 'N/A'),
        ('Alcohol', 'Duration', 'N/A'),
        ('Tobacco', 'StatusTime', 'current'),
        ('Tobacco', 'StatusTime', 'none'),
        ('Tobacco', 'StatusTime', 'past'),
        ('Tobacco', 'Type', 'N/A'),
        ('Tobacco', 'Frequency', 'N/A'),
        ('Tobacco', 'Amount', 'N/A'),
        ('Tobacco', 'History', 'N/A'),
        ('Tobacco', 'Duration', 'N/A'),
        ('Drug', 'StatusTime', 'current'),
        ('Drug', 'StatusTime', 'none'),
        ('Drug', 'StatusTime', 'past'),
        ('Drug', 'Type', 'N/A'),
        ('Drug', 'Frequency', 'N/A'),
        ('Drug', 'Amount', 'N/A'),
        ('Drug', 'History', 'N/A'),
        ('Drug', 'Duration', 'N/A'),
        ('PhysicalActivity_Yes', 'N/A', 'N/A'),
        ('PhysicalActivity_Yes', 'Frequency', 'N/A'),
        ('PhysicalActivity_Yes', 'Amount', 'N/A'),
        ('PhysicalActivity_Yes', 'Duration', 'N/A'),
        ('PhysicalActivity_No', 'History', 'N/A'),
        ('Ethnicity', 'N/A', 'N/A'),
        ('Education', 'N/A', 'N/A'),
        ('Education', 'History', 'N/A'),
        ('Income', 'N/A', 'N/A'),
        ('Income', 'History', 'N/A')
    ]

    updated_results = []

    for entry in json_entries:
        doc_id = entry['document_id']
        text = entry['fr']

        entry['gold_format'] = format_struct_output(entry['sdoh'])
        entry['pred_format'] = format_struct_output(entry['sdoh_generated'])

        # Update offsets in predictions
        pred_updated = []
        for e in entry['pred_format']:
            e = update_entity_offsets(e, text, pred_updated)
            if e:
                e = update_relation_offsets(e, text)
                e = update_status_relations(e)
                pred_updated.append(e)

        # Load and parse gold annotation
        brat_path = Path(args.brat_dir) / f"{doc_id}.txt"
        gold_brat = parse_brat_file(brat_path, parse_notes=True)
        gold_struct = convert_text_to_struct(gold_brat)

        entry['gold_format_updated'] = gold_struct
        entry['pred_format_updated'] = pred_updated
        updated_results.append(entry)

    # Run evaluation
    df = build_eval_dataframe(updated_results, possible_combinations, eval_type=args.eval_type)
    print(df)

    # Save to Word
    save_df_to_word(df, args.output)


if __name__ == '__main__':
    main()
